import re
from pathlib import Path
from typing import Optional, Dict, List, Set
from app.services.resolver import SkillResolver

TECHNICAL_SKILLS = {
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'csharp', 'ruby', 'go', 'golang',
    'rust', 'php', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'haskell', 'elixir', 'clojure',
    'sql', 'html', 'css', 'sass', 'scss', 'less', 'xml', 'json', 'yaml', 'toml',
    
    # Frameworks & Libraries
    'react', 'reactjs', 'vue', 'vuejs', 'angular', 'angularjs', 'svelte', 'nextjs', 'nuxt',
    'nodejs', 'node.js', 'express', 'expressjs', 'fastapi', 'django', 'flask', 'rails', 'spring',
    'spring boot', 'asp.net', 'laravel', 'phoenix', 'gin', 'echo', 'fiber',
    
    # Frontend
    'tailwind', 'bootstrap', 'material ui', 'mui', 'chakra ui', 'styled-components',
    'webpack', 'vite', 'esbuild', 'rollup', 'babel', 'eslint', 'prettier',
    
    # Databases
    'postgresql', 'postgres', 'mysql', 'mariadb', 'mongodb', 'mongo', 'redis', 'elasticsearch',
    'elastic', 'cassandra', 'dynamodb', 'sqlite', 'sql server', 'oracle', 'snowflake', 'bigquery',
    'neo4j', 'couchdb', 'firebase', 'supabase',
    
    # Cloud & DevOps
    'aws', 'amazon web services', 'azure', 'microsoft azure', 'gcp', 'google cloud',
    'docker', 'kubernetes', 'k8s', 'helm', 'terraform', 'ansible', 'puppet', 'chef',
    'jenkins', 'gitlab ci', 'github actions', 'circleci', 'travis', 'bitbucket pipelines',
    
    # Cloud Services
    'ec2', 's3', 'lambda', 'ecs', 'eks', 'eks', 'gke', 'aks', 'rds', 'dynamodb',
    'cloudformation', 'cdk', 'route53', 'cloudwatch', 'sns', 'sqs', 'api gateway',
    'vpc', 'iam', 'cloudtrail', 'kinesis', 'glue', 'athena', 'redshift',
    
    # AI/ML
    'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'sklearn', 'pandas', 'numpy', 'scipy',
    'machine learning', 'deep learning', 'neural network', 'nlp', 'computer vision',
    'ai', 'artificial intelligence', 'data science', 'ml', 'llm', 'openai', 'hugging face',
    'langchain', 'opencv', 'spacy', 'nltk', 'xgboost', 'lightgbm', 'catboost',
    
    # Tools & Version Control
    'git', 'github', 'gitlab', 'bitbucket', 'svn', 'mercurial', 'jira', 'confluence',
    'notion', 'asana', 'trello', 'slack', 'datadog', 'new relic', 'grafana', 'prometheus',
    
    # API & Integration
    'rest', 'restful', 'graphql', 'grpc', 'websocket', 'webhook', 'api', 'microservices',
    'oauth', 'jwt', 'oauth2', 'saml', 'openid connect',
    
    # Testing
    'testing', 'tdd', 'bdd', 'unit test', 'integration test', 'e2e', 'selenium', 'cypress',
    'playwright', 'jest', 'mocha', 'pytest', 'junit', 'testng', 'pytest',
    
    # Data Engineering
    'spark', 'hadoop', 'hive', 'airflow', 'dbt', 'etl', 'data pipeline', 'kafka',
    'rabbitmq', 'flink', 'storm', 'beam', 'databricks',
    
    # Operating Systems
    'linux', 'unix', 'windows', 'macos', 'ubuntu', 'centos', 'debian', 'redhat',
    'bash', 'shell', 'zsh', 'powershell', 'windows server',
    
    # Methodologies
    'agile', 'scrum', 'kanban', 'waterfall', 'devops', 'sre', 'clean code',
    
    # Security
    'security', 'cybersecurity', 'ssl', 'tls', 'https', 'encryption', 'oauth', 'jwt',
    'owasp', 'penetration testing', 'vulnerability assessment', 'security audit',
    
    # Mobile
    'react native', 'flutter', 'swift', 'kotlin', 'ios', 'android', 'xamarin',
    
    # Other
    'nginx', 'apache', 'caddy', 'traefik', 'envoy', 'kong', 'rabbitmq',
    'wordpress', 'drupal', 'shopify', 'magento',
}

SKILL_VARIATIONS = {
    'js': 'javascript',
    'ts': 'typescript',
    'tsx': 'typescript',
    'jsx': 'javascript',
    'py': 'python',
    'rb': 'ruby',
    'go': 'golang',
    'k8s': 'kubernetes',
    'mongo': 'mongodb',
    'postgres': 'postgresql',
    'pg': 'postgresql',
    'eks': 'kubernetes',
    'gke': 'kubernetes',
    'aks': 'kubernetes',
    'tf': 'terraform',
    'kafka': 'kafka',
    'ml': 'machine learning',
    'dl': 'deep learning',
    'ai': 'artificial intelligence',
    'ds': 'data science',
    'de': 'data engineering',
    'fe': 'frontend',
    'be': 'backend',
    'fs': 'fullstack',
    'react.js': 'react',
    'vue.js': 'vue',
    'node.js': 'nodejs',
    'next.js': 'nextjs',
    'aws': 'aws',
    'gcp': 'google cloud',
}


class ResumeParser:
    def __init__(self):
        self.resolver = SkillResolver()
        self.tech_skills_set = TECHNICAL_SKILLS
        self.skill_variations = SKILL_VARIATIONS
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
            
            combined = "\n".join(text_parts)
            return self._clean_extracted_text(combined)
            
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    def _clean_extracted_text(self, text: str) -> str:
        if not text:
            return ""
        
        import re
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Try to fix column layouts (common in resumes)
        lines = text.split('\n')
        fixed_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                fixed_lines.append('')
                continue
            
            # If line has pipe separators, split it
            if '|' in line and len(line) > 60:
                parts = line.split('|')
                for part in parts:
                    part = part.strip()
                    if part:
                        fixed_lines.append(part)
            else:
                fixed_lines.append(line)
        
        return '\n'.join(fixed_lines)
    
    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append(' | '.join(row_texts))
            
            combined = '\n'.join(text_parts)
            return self._clean_extracted_text(combined)
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
    
    def extract_text(self, file_path: str) -> str:
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif extension == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                return f"Error reading TXT: {str(e)}"
        else:
            return f"Unsupported file format: {extension}"
    
    def extract_skills_from_text(self, text: str) -> List[str]:
        found_skills: Set[str] = set()
        text_lower = text.lower()
        
        all_patterns = list(self.tech_skills_set) + list(self.skill_variations.keys())
        
        for pattern_str in all_patterns:
            pattern_str_lower = pattern_str.lower()
            pattern = r'\b' + re.escape(pattern_str_lower) + r'\b'
            if re.search(pattern, text_lower):
                if pattern_str_lower in self.skill_variations:
                    canonical = self.skill_variations[pattern_str_lower].replace(' ', '_')
                    found_skills.add(canonical)
                else:
                    found_skills.add(pattern_str_lower.replace(' ', '_'))
        
        found_skills.discard('')
        found_skills.discard(' ')
        
        return list(found_skills)
    
    def extract_skills_from_skill_section(self, skill_text: str) -> List[str]:
        found_skills: Set[str] = set()
        text_lower = skill_text.lower()
        
        for skill in self.tech_skills_set:
            skill_lower = skill.lower()
            pattern = r'\b' + re.escape(skill_lower) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill_lower.replace(' ', '_'))
        
        for variation, canonical in self.skill_variations.items():
            pattern = r'\b' + re.escape(variation) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(canonical.replace(' ', '_'))
        
        found_skills.discard('')
        found_skills.discard(' ')
        
        return list(found_skills)
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        sections = {
            "header": "",
            "summary": "",
            "experience": "",
            "education": "",
            "skills": "",
            "certifications": "",
            "projects": "",
            "other": ""
        }
        
        lines = text.split('\n')
        current_section = "other"
        header_pattern = r'^[A-Z][A-Z\s]+$|Contact|Phone|Email|LinkedIn|Portfolio'
        
        section_keywords = {
            "summary": ["summary", "objective", "profile", "about me", "professional summary"],
            "experience": ["experience", "employment", "work history", "professional experience", "work experience", "career history"],
            "education": ["education", "academic", "degree", "university", "college", "school", "certification"],
            "skills": ["skills", "technical skills", "technologies", "competencies", "expertise", "technical proficiency", "core competencies", "technical competencies"],
            "certifications": ["certification", "certifications", "credentials", "licensed", "professional certifications"],
            "projects": ["projects", "portfolio", "personal projects", "open source", "github", "contributions"],
        }
        
        # Remove markdown-style headers first
        text = re.sub(r'===([^=]+)===\s*', r'\n## \1 ##\n', text)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        
        lines = text.split('\n')
        current_section = "other"
        
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue
            
            line_lower = line_clean.lower()
            
            # Check for section headers (## headers or standalone keywords)
            matched_section = None
            for section, keywords in section_keywords.items():
                for kw in keywords:
                    if line_lower == kw or line_lower == f'## {kw} ##' or line_lower.startswith(f'{kw} :') or line_lower.startswith(f'{kw}:'):
                        matched_section = section
                        break
                if matched_section:
                    break
            
            if matched_section:
                current_section = matched_section
            elif current_section == "other":
                sections["header"] += line_clean + '\n'
            else:
                sections[current_section] += line_clean + '\n'
        
        for key in sections:
            sections[key] = sections[key].strip()
        
        return sections
    
    def extract_name(self, text: str) -> Optional[str]:
        lines = text.split('\n')
        for line in lines[:5]:
            line_clean = line.strip()
            if line_clean and len(line_clean) < 50 and len(line_clean.split()) <= 4:
                if not any(char.isdigit() for char in line_clean):
                    if re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)+$', line_clean):
                        return line_clean
        return None
    
    def extract_contact(self, text: str) -> Dict[str, str]:
        contact = {}
        
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact["email"] = emails[0]
        
        phone_pattern = r'(\+?1?\s*[-.]?\s*)?(\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            parts = phones[0]
            phone_str = ''
            for part in parts:
                if part:
                    phone_str += ''.join(filter(str.isdigit, part))
            if phone_str:
                contact["phone"] = phone_str
        
        linkedin_pattern = r'linkedin\.com/in/([a-zA-Z0-9-]+)'
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            contact["linkedin"] = linkedin[0]
        
        github_pattern = r'github\.com/([a-zA-Z0-9-]+)'
        github = re.findall(github_pattern, text, re.IGNORECASE)
        if github:
            contact["github"] = github[0]
        
        return contact
    
    def parse_resume(self, file_path: str) -> Dict:
        text = self.extract_text(file_path)
        
        if text.startswith("Error"):
            return {
                "success": False,
                "error": text
            }
        
        if len(text.strip()) < 50:
            return {
                "success": False,
                "error": "Could not extract enough text from resume. The file may be scanned or image-based."
            }
        
        sections = self.extract_sections(text)
        
        # Extract skills from dedicated skills section (highest priority)
        skills_section_skills = self.extract_skills_from_skill_section(sections.get("skills", ""))
        
        # Extract skills from experience section
        experience_skills = self.extract_skills_from_text(sections.get("experience", ""))
        
        # Extract skills from projects section
        project_skills = self.extract_skills_from_text(sections.get("projects", ""))
        
        # Extract skills from certifications section
        cert_skills = self.extract_skills_from_text(sections.get("certifications", ""))
        
        # Extract skills from summary section
        summary_skills = self.extract_skills_from_text(sections.get("summary", ""))
        
        # Extract skills from header (contact info often contains links to profiles)
        header_skills = self.extract_skills_from_text(sections.get("header", ""))
        
        # Also extract from the full text as fallback
        all_skills = self.extract_skills_from_text(text)
        
        # Combine all found skills (deduplicated)
        all_found_skills = list(set(
            skills_section_skills + 
            experience_skills + 
            project_skills + 
            cert_skills + 
            summary_skills + 
            header_skills +
            all_skills
        ))
        
        # Resolve skills using pattern matching (fast, no LLM needed)
        resolved_skills = list(set(self._quick_resolve(all_found_skills)))
        resolved_skills = [s for s in resolved_skills if s]
        
        return {
            "success": True,
            "name": self.extract_name(text),
            "contact": self.extract_contact(text),
            "sections": sections,
            "raw_skills": all_found_skills,
            "resolved_skills": resolved_skills,
            "summary": sections.get("summary", "")[:500],
            "experience_years": self._estimate_years(sections.get("experience", ""))
        }
    
    async def parse_resume_llm(self, file_path: str) -> Dict:
        text = self.extract_text(file_path)
        
        if text.startswith("Error"):
            return {
                "success": False,
                "error": text
            }
        
        if len(text.strip()) < 50:
            return {
                "success": False,
                "error": "Could not extract enough text from resume."
            }
        
        sections = self.extract_sections(text)
        
        skills_section_skills = self.extract_skills_from_skill_section(sections.get("skills", ""))
        experience_skills = self.extract_skills_from_text(sections.get("experience", ""))
        project_skills = self.extract_skills_from_text(sections.get("projects", ""))
        cert_skills = self.extract_skills_from_text(sections.get("certifications", ""))
        summary_skills = self.extract_skills_from_text(sections.get("summary", ""))
        header_skills = self.extract_skills_from_text(sections.get("header", ""))
        all_skills = self.extract_skills_from_text(text)
        
        all_found_skills = list(set(
            skills_section_skills + 
            experience_skills + 
            project_skills + 
            cert_skills + 
            summary_skills + 
            header_skills +
            all_skills
        ))
        
        # Use simple resolution (fast, no LLM)
        resolved_skills = list(set(self._quick_resolve(all_found_skills)))
        resolved_skills = list(set(resolved_skills.values()))
        resolved_skills = [s for s in resolved_skills if s]
        
        return {
            "success": True,
            "name": self.extract_name(text),
            "contact": self.extract_contact(text),
            "sections": sections,
            "raw_skills": all_found_skills,
            "resolved_skills": resolved_skills,
            "summary": sections.get("summary", "")[:500],
            "experience_years": self._estimate_years(sections.get("experience", "")),
            "resolution_method": "llm"
        }
    
    def _estimate_years(self, experience_text: str) -> Optional[int]:
        year_pattern = r'(?:19|20)\d{2}'
        years = re.findall(year_pattern, experience_text)
        if len(years) >= 2:
            years_int = [int(y) for y in years]
            return max(years_int) - min(years_int)
        return None
    
    def _quick_resolve(self, skills: List[str]) -> List[str]:
        """Quick skill resolution without LLM - uses pattern matching"""
        resolved = []
        for skill in skills:
            skill_lower = skill.lower()
            if skill_lower in self.skill_variations:
                resolved.append(self.skill_variations[skill_lower].replace(' ', '_'))
            else:
                resolved.append(skill_lower.replace(' ', '_'))
        return resolved
    
    def parse_resume_bytes(self, content: bytes, filename: str, temp_dir: str = None) -> Dict:
        import tempfile
        import os
        
        suffix = Path(filename).suffix.lower()
        
        if temp_dir is None:
            temp_dir = tempfile.gettempdir()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=temp_dir) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            return self.parse_resume(tmp_path)
        finally:
            try:
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
            except Exception:
                pass
